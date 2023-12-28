import io
import os
import gradio as gr
import json
import base64
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

os.environ["no_proxy"] = "localhost,127.0.0.1,::1"

# h1_counter = 0
# h2_counter = 1  #1.
# h3_counter = 0  #가.
# h4_counter = 1  #1)
# h5_counter = 0  #가)
# h6_counter = 1  #(1)
# h7_counter = 0  #(가)
counters = [0, 1, 0, 1, 0, 1, 0]
h1_headers = ['I. ', 'II. ', 'III. ', 'IV. ', 'V. ', 'VI. ', 'VII. ', 'VIII. ', 'IX. ', 'X. ']
korean_headers = ['가', '나', '다', '라', '마', '바', '사', '아', '자', '차', '카', '타', '파', '하']
previous_level = 0
image_counter = 1
table_counter = 1
ref_counter = 1
doc = Document('FinalReport_empty.docx')
def get_current_header(level, counter):
    if level == 1:
        return str(counter) + '. '
    elif level == 2:
        return korean_headers[counter] + '. '
    elif level == 3:
        return str(counter) + ') '
    elif level == 4:
        return korean_headers[counter] + ') '
    elif level == 5:
        return '(' + str(counter) + ') '
    elif level == 6:
        return '(' + korean_headers[counter] + ') '


def generate_metadata(file_name, school_name, student_names, field_name, mentor_name, start_date, end_date, title, title_eng, keywords):
    global doc
    doc.tables[0].rows[0].cells[2].paragraphs[0].text = school_name
    doc.tables[0].rows[0].cells[4].paragraphs[0].text = student_names
    doc.tables[0].rows[1].cells[2].paragraphs[0].text = field_name
    doc.tables[0].rows[1].cells[4].paragraphs[0].text = mentor_name
    doc.tables[0].rows[2].cells[1].paragraphs[0].text = f"{start_date[0:4]}년 {start_date[4:6]}월 {start_date[6:]}일 ~ {end_date[0:4]}년 {end_date[4:6]}월 {end_date[6:]}일"
    doc.tables[0].rows[3].cells[2].paragraphs[0].text = title
    doc.tables[0].rows[4].cells[2].paragraphs[0].text = title_eng
    doc.tables[0].rows[5].cells[1].paragraphs[0].text = keywords
    doc.save(file_name + '.docx')
    doc = Document('FinalReport_empty.docx')
    print("Done")

def parse_json(file):
    global previous_level, image_counter, ref_counter, counters, doc
    with open(file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    print(data)
    for content in data['contents']:
        if content['type'] == 'h1':
            current_level = int(content['type'][1]) - 1
            if current_level < previous_level:
                for i in range(current_level + 1, 7):
                    counters[i] = i % 2

            para = doc.tables[1].rows[0].cells[0].add_paragraph()
            run = para.add_run(h1_headers[counters[current_level]] + content['content'])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '맑은 고딕'
            counters[current_level] += 1
            previous_level = current_level

        elif content['type'] == 'text':
            para = doc.tables[1].rows[0].cells[0].add_paragraph()
            run = para.add_run(' ' + content['content'] + '\n')
            run.bold = False
            run.font.size = Pt(10)

        elif content['type'] == 'image':
            para = doc.tables[1].rows[0].cells[0].add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            binary = base64.b64decode(content['content'].split(',')[1])
            run.add_picture(io.BytesIO(binary), height=Inches(3))
            run2 = para.add_run('\n그림 ' + str(image_counter) + '. ' + content['caption'] + '\n')
            run2.font.size = Pt(8)
            image_counter += 1

        elif content['type'] == "ref":
            para = doc.tables[1].rows[0].cells[0].add_paragraph()
            run = para.add_run(f'[{ref_counter}] ' + content['content'])
            run.bold = False
            run.font.size = Pt(10)
            ref_counter += 1

        else:
            current_level = int(content['type'][1]) - 1
            if current_level < previous_level:
                for i in range(current_level + 1, 7):
                    counters[i] = i % 2

            para = doc.tables[1].rows[0].cells[0].add_paragraph()
            para.paragraph_format.left_indent = Pt(20 * current_level)
            run = para.add_run(get_current_header(current_level, counters[current_level]) + content['content'])
            run.bold = False
            run.font.size = Pt(10)
            counters[current_level] += 1
            previous_level = current_level

    image_counter = 1
    ref_counter = 1
    counters = [0, 1, 0, 1, 0, 1, 0]
    previous_level = 0
    return

with gr.Blocks() as demo:
    file_name = gr.Textbox(label="파일 제목")
    with gr.Column():
        school_name = gr.Textbox(label="학교명")
        student_names = gr.Textbox(label="학번/이름")
    with gr.Column():
        field_name = gr.Textbox(label="분야")
        mentor_name = gr.Textbox(label="멘토명")
    with gr.Column():
        start_date = gr.Textbox(label="시작 날짜 (Ex. YYYYMMDD)")
        end_date = gr.Textbox(label="종료 날짜 (Ex. YYYYMMDD)")
    title = gr.Textbox(label="연구 제목")
    title_eng = gr.Textbox(label="연구 제목(영어)")
    keywords = gr.Textbox(label="키워드 (5개 이상)")
    submit_button = gr.Button(value="Submit")

    file = gr.File(label='JSON 파일', file_types=['.json'], file_count='single')
    file.upload(parse_json, file)
    submit_button.click(fn=generate_metadata, inputs=[file_name,
                                                      school_name,
                                                      student_names,
                                                      field_name,
                                                      mentor_name,
                                                      start_date,
                                                      end_date,
                                                      title,
                                                      title_eng,
                                                      keywords,
                                                      ])


demo.launch()