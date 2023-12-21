from docx import Document
from docx.text.font import Font

doc = Document('FinalReport.docx')

# for i, table in enumerate(doc.tables):
#     for j, row in enumerate(table.rows):
#         for k, cell in enumerate(row.cells):
#             for p, para in enumerate(cell.paragraphs):
#                 print(f"Table {i} / row {j} / col {k} / para {p}")
#                 print(para.text)

doc.tables[0].rows[0].cells[2].paragraphs[0].text = "한일고등학교"
doc.tables[0].rows[0].cells[4].paragraphs[0].text = "3425 임재석"
doc.tables[0].rows[1].cells[2].paragraphs[0].text = "빅데이터 정책제안 프로젝트"
doc.tables[0].rows[1].cells[4].paragraphs[0].text = "정영우 멘토"
doc.tables[0].rows[2].cells[1].paragraphs[0].text = "2023년 11월 23일 ~ 2023년 11월 24일"
doc.tables[0].rows[3].cells[2].paragraphs[0].text = "화가 나게 하는 인공지능 화학"
doc.tables[0].rows[4].cells[2].paragraphs[0].text = "Annoying AI Chemistry"
doc.tables[0].rows[5].cells[1].paragraphs[0].text = "화학, 인공지능, 분노, 분조장, 억까"

doc.save("rev_FinalReport.docx")