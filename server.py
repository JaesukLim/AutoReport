from flask import Flask, render_template, url_for, redirect, request, jsonify, send_file, g
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
import json
import os
import math
import io
import base64
import shutil
import glob

# For Exporting Final Report
counters = [0, 1, 0, 1, 0, 1, 0]
h1_headers = ['I. ', 'II. ', 'III. ', 'IV. ', 'V. ', 'VI. ', 'VII. ', 'VIII. ', 'IX. ', 'X. ']
korean_headers = ['가', '나', '다', '라', '마', '바', '사', '아', '자', '차', '카', '타', '파', '하']
previous_level = 0
image_counter = 1
table_counter = 1
ref_counter = 1

def get_current_header(level, counter):
    if level == 1:
        return str(counter) + '. '
    elif level == 2:
        return korean_headers[counter] + '. '
    elif level == 3:
        return str(counter) + ') '
    elif level == 4:
        return korean_headers[counter] + ') '
    elif level == 5:
        return '(' + str(counter) + ') '
    elif level == 6:
        return '(' + korean_headers[counter] + ') '

app = Flask(__name__)
@app.route('/admin')
def mentor_login():
    return render_template('user.html')

@app.route('/dashboard/<mentor_uuid>')
def mentor_dashboard(mentor_uuid):
    return render_template('dashboard.html')

@app.route('/login', methods=['POST'])
def login():
    data = request.json
    with open('./assets/mentor_ids.json', 'r') as f:
        ids = json.load(f)

    if data['nickname'] not in list(ids['id_password'].keys()) or data["password"] != ids["id_password"][data["nickname"]]:
        return jsonify({"message": "비밀번호가 틀렸습니다"})

    else:
        return redirect(url_for('mentor_dashboard', mentor_uuid=ids["id_uuid"][data["nickname"]]))

@app.route('/register', methods=['POST'])
def register():
    data = request.form
    with open('./assets/mentor_ids.json', 'r') as f:
        ids = json.load(f)

    if data['nickname'] == "":
        return jsonify({"message" : "이름을 입력해주세요"})

    if data['nickname'] in list(ids['id_password'].keys()):
        return jsonify({"message": "이미 등록된 이름입니다"})

    if data['password'] != data['passwordCheck']:
        return jsonify({"message": "두 비밀번호가 다릅니다"})

    #JSON에 멘토명-비밀번호 등록
    ids["id_password"][data["nickname"]] = data["password"]
    temp_uuid = str(uuid.uuid4())[:8]
    ids["id_uuid"][data["nickname"]] = temp_uuid
    ids["uuid_id"][temp_uuid] = data["nickname"]
    with open('./assets/mentor_ids.json', 'w') as f:
        json.dump(ids, f)
    #멘토명 폴더 생성
    os.mkdir(os.path.join('./assets', temp_uuid))
    with open(os.path.join('./assets', temp_uuid, 'school_ids.json'), 'w') as f:
        json.dump({"school_uuid" : {}, "uuid_school" : {}}, f)
    return jsonify({"message" : "성공적으로 등록되었습니다"})

@app.route('/school/<mentor_uuid>', methods=['POST', 'GET'])
def add_school_and_student(mentor_uuid):
    if request.method == 'POST':
        data = request.form
        with open('./assets/mentor_ids.json', 'r') as f:
            ids = json.load(f)
        school_dir = os.path.join('assets', mentor_uuid, data["school-name"] + ' ' + data["field"])
        os.mkdir(school_dir)
        students = []
        report = {}
        student_num_uuid = {}
        uuid_student_num = {}
        for i in range(1, ((len(data) - 2) // 2) + 1):
            students.append({data["student-number-" + str(i)] : data["student-name-" + str(i)]})
            report[data["student-number-" + str(i)]] = {"registered": "No", "final_report": "No", "self_evaluation": "No", "self_evaluation_revised": "No", "password": ""}
            temp_uuid = str(uuid.uuid4())[:8]
            student_num_uuid[data["student-number-" + str(i)]] = temp_uuid
            uuid_student_num[temp_uuid] = data["student-number-" + str(i)]

        metadata = {
            "school_name" : data["school-name"],
            "field" : data["field"],
            "students" : students,
            "class_info" : {},
            "attendance" : {},
            "class_screenshot" : {},
            "report_status" : report,
            "student_num_uuid" : student_num_uuid,
            "uuid_student_num" : uuid_student_num
        }
        with open(os.path.join(school_dir, "metadata.json"), 'w') as f:
            json.dump(metadata, f)
        with open(os.path.join('assets', mentor_uuid, 'school_ids.json'), 'r') as f:
            school_ids = json.load(f)
        cur_school_id = str(uuid.uuid4())[:8]
        school_ids["school_uuid"][data["school-name"] + ' ' + data['field']] = cur_school_id
        school_ids["uuid_school"][cur_school_id] = data["school-name"] + ' ' + data['field']
        with open(os.path.join('assets', mentor_uuid, 'school_ids.json'), 'w') as f:
            json.dump(school_ids, f)
        return "YAAY"

    elif request.method == 'GET':
        res = os.listdir(os.path.join('./assets', mentor_uuid))
        res.remove('school_ids.json')
        return res

@app.route('/students/<mentor_uuid>', methods=["POST"])
def get_students(mentor_uuid):
    data = request.json
    school_name = data["school_name"]
    with open(os.path.join('assets', mentor_uuid, school_name, 'metadata.json'), 'r') as f:
        school_info = json.load(f)
    students = school_info["students"]
    return students

@app.route('/metadata/<mentor_uuid>', methods=['POST'])
def get_class_metadata(mentor_uuid):
    data = request.form
    print(data)
    school_dir = os.path.join('assets', mentor_uuid, data['schoolName'], 'metadata.json')
    with open(school_dir, 'r') as f:
        school_metadata = json.load(f)
    school_metadata['class_info'] = data
    with open(school_dir, 'w') as f:
        json.dump(school_metadata, f)
    return jsonify({"message" : "Success"})

@app.route('/attendance/<mentor_uuid>', methods=['POST'])
def get_attendance(mentor_uuid):
    data = request.form
    print(data)
    school_dir = os.path.join('assets', mentor_uuid, data['schoolName'], 'metadata.json')
    with open(school_dir, 'r') as f:
        school_metadata = json.load(f)
    school_metadata['attendance'] = data
    with open(school_dir, 'w') as f:
        json.dump(school_metadata, f)
    return jsonify({"message" : "Success"})

@app.route('/classdata/<mentor_uuid>', methods=['POST'])
def get_class_data(mentor_uuid):
    data = request.json
    with open(os.path.join('assets', mentor_uuid, data['school_name'], 'metadata.json'), 'r') as f:
        metadata = json.load(f)
    with open(os.path.join('assets', mentor_uuid, 'school_ids.json'), 'r') as f:
        school_ids = json.load(f)
    metadata['school_id'] = school_ids['school_uuid'][data['school_name']]
    return jsonify(metadata)

@app.route('/screenshots/<mentor_uuid>', methods=['POST'])
def get_screenshots(mentor_uuid):
    data = request.json
    print(data)
    school_dir = os.path.join('assets', mentor_uuid, data['school_name'], 'metadata.json')
    with open(school_dir, 'r') as f:
        metadata = json.load(f)
    data.pop('school_name')
    metadata['class_screenshot'] = data
    with open(school_dir, 'w') as f:
        json.dump(metadata, f)
    return jsonify({"message" : "Success"})

@app.route('/export/<mentor_uuid>', methods=['POST'])
def get_class_report(mentor_uuid):
    data = request.form
    for file in os.listdir(os.path.join('assets', mentor_uuid, data['school_name'])):
        if '.docx' in file:
            os.remove(os.path.join('assets', mentor_uuid, data['school_name'], file))
    school_dir = os.path.join('assets', mentor_uuid, data['school_name'], 'metadata.json')
    with open(school_dir, 'r') as f:
        metadata = json.load(f)
    doc = Document('ClassReport_Empty.docx')
    doc.tables[0].rows[0].cells[2].paragraphs[0].text = ''
    run = doc.tables[0].rows[0].cells[2].paragraphs[0].add_run(metadata['school_name'])
    run.font.size = Pt(12)
    doc.tables[0].rows[1].cells[2].paragraphs[0].text = ''
    run = doc.tables[0].rows[1].cells[2].paragraphs[0].add_run('  ' + metadata['field'] + ' ' + metadata['class_info']['className'])
    run.font.size = Pt(12)
    doc.tables[0].rows[2].cells[2].paragraphs[0].text = ''
    run = doc.tables[0].rows[2].cells[2].paragraphs[0].add_run(metadata['class_info']['mentorName'])
    run.font.size = Pt(12)
    for i in range(1, 9):
        doc.tables[1].rows[i].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.tables[1].rows[i].cells[1].paragraphs[0].text = metadata['class_info']['classOverview' + str(i)].replace('\r\n', '\n')
        doc.tables[1].rows[i].cells[2].paragraphs[0].text = metadata['class_info']['classDetails' + str(i)].replace('\r\n', '\n')
    doc.tables[1].rows[9].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.tables[1].rows[9].cells[1].paragraphs[0].text = metadata['class_info']['classOverviewOffline'].replace('\r\n', '\n')
    doc.tables[1].rows[9].cells[2].paragraphs[0].text = metadata['class_info']['classDetailsOffline'].replace('\r\n', '\n')

    doc.tables[2].rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.tables[2].rows[1].cells[1].paragraphs[0].text = metadata['class_info']['teamTopic'].replace('\r\n', '\n')
    doc.tables[2].rows[2].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.tables[2].rows[2].cells[1].paragraphs[0].text = metadata['class_info']['individualInterest'].replace('\r\n', '\n')

    cur_row = 2
    for student in metadata['students']:
        student_number = list(student.keys())[0]
        student_name = student[student_number]
        doc.tables[3].rows[cur_row].cells[0].paragraphs[0].text = student_number
        doc.tables[3].rows[cur_row].cells[1].paragraphs[0].text = student_name
        for key in metadata['attendance'].keys():
            status_str = ""
            if 'attendance-' + student_number in key:
                col_num = int(key.split('-')[-1])
                if metadata['attendance'][key] == '출석':
                    status_str = 'O'
                elif metadata['attendance'][key] == '결석':
                    status_str = 'X'
                elif metadata['attendance'][key] == '지각':
                    status_str = metadata['attendance'][key]

                if status_str == 'X' or status_str == '지각':
                    status_str += '\n(' + metadata['attendance']['reason-' + student_number + '-' + str(col_num)] + ')'
                doc.tables[3].rows[cur_row].cells[col_num + 1].paragraphs[0].text = status_str
            elif 'time-' in key:
                col_num = int(key.split('-')[-1])
                doc.tables[3].rows[1].cells[col_num + 1].paragraphs[0].text = ''
                run = doc.tables[3].rows[1].cells[col_num + 1].paragraphs[0].add_run(metadata['attendance'][key].replace('\r\n', '\n'))
                run.font.size = Pt(9)
        cur_row += 1

    for k, v in metadata['class_screenshot'].items():
        binary = base64.b64decode(v.split(',')[1])
        run = doc.tables[3 + math.ceil(int(k) / 2)].rows[1].cells[(int(k) + 1) % 2].paragraphs[0].add_run()
        run.add_picture(io.BytesIO(binary), width=Inches(3))

    final_dir = os.path.join('assets', mentor_uuid, data['school_name'], data['file_name'] + '.docx')
    doc.save(final_dir)
    return send_file(final_dir, download_name=data['file_name'] + '.docx', as_attachment=True)

@app.route('/student_login', methods=['GET', 'POST'])
def student_login():
    if request.method == 'GET':
        args = request.args
        mentor_code = args.get('mentor_code', type=str)
        school_code = args.get('school_code', type=str)
        return render_template('user_student.html', mentor_code=mentor_code, school_code=school_code)
    elif request.method == 'POST':
        data = request.json
        args = request.args
        print(data)
        mentor_uuid = args['mentor_code']
        school_uuid = args['school_code']
        with open(os.path.join('assets', mentor_uuid, 'school_ids.json'), 'r') as f:
            school_ids = json.load(f)
        school_name = school_ids["uuid_school"][school_uuid]
        with open(os.path.join('assets', mentor_uuid, school_name, 'metadata.json'), 'r') as f:
            metadata = json.load(f)

        student_number = ''
        for student in metadata['students']:
            if list(student.values())[0] == data['nickname']:
                student_number = list(student.keys())[0]

        print("student num:", student_number)
        if student_number == '':
            return jsonify({"message" : "학번을 찾을 수 없습니다"})

        if metadata['report_status'][student_number]['registered'] == 'No':
            return jsonify({"message" : "등록되지 않은 학생입니다"})

        if data['password'] != metadata['report_status'][student_number]['password']:
            return jsonify({"message" : "비밀번호가 틀렸습니다"})

        else:
            return redirect(url_for('report_page',
                                    mentor_code=mentor_uuid,
                                    school_code=school_uuid,
                                    student_code=metadata['student_num_uuid'][student_number]))


@app.route('/register_student', methods=['POST'])
def register_student():
    data = request.form
    args = request.args
    print(data)
    mentor_uuid = args['mentor_code']
    school_uuid = args['school_code']
    with open(os.path.join('assets', mentor_uuid, 'school_ids.json'), 'r') as f:
        school_ids = json.load(f)
    school_name = school_ids["uuid_school"][school_uuid]
    with open(os.path.join('assets', mentor_uuid, school_name, 'metadata.json'), 'r') as f:
        metadata = json.load(f)

    flag = False
    student_name = ''
    for elem in metadata['students']:
        if data['studentNumber'] == list(elem.keys())[0]:
            flag = True
            student_name = list(elem.values())[0]
    if not flag:
        return jsonify({"message" : "학번이 존재하지 않습니다", 'status': "Fail"})

    if student_name != '' and data['studentName'] != student_name:
        return jsonify({"message" : "학번과 이름이 다릅니다", 'status': "Fail"})

    if data['password'] != data['passwordCheck']:
        return jsonify({"message" : "두 비밀번호가 다릅니다", 'status': "Fail"})

    if metadata['report_status'][data['studentNumber']]['registered'] == 'Yes':
        return jsonify({"message" : "이미 등록된 학생입니다", 'status': "Fail"})

    metadata['report_status'][data['studentNumber']]['registered'] = 'Yes'
    metadata['report_status'][data['studentNumber']]['password'] = data['password']
    with open(os.path.join('assets', mentor_uuid, school_name, 'metadata.json'), 'w') as f:
        json.dump(metadata, f)
    return jsonify({"message" : "성공적으로 등록되었습니다\n만든 계정으로 로그인하세요", 'status': "Success"})

@app.route('/report_page')
def report_page():
    return render_template('main.html')

@app.route('/student_id/<mentor_uuid>', methods=['POST'])
def get_student_id(mentor_uuid):
    data = request.json
    school_name = data['school_name']
    student_number = data['student_number']
    with open(os.path.join('assets', mentor_uuid, school_name, 'metadata.json'), 'r') as f:
        metadata = json.load(f)
    return jsonify({"student_id" : metadata['student_num_uuid'][student_number]})



@app.route('/final_report', methods=['GET', 'POST'])
def final_report():
    if request.method == 'GET':
        mentor_uuid = request.args['mentor_code']
        school_uuid = request.args['school_code']
        student_uuid = request.args['student_code']
        with open(os.path.join('assets', mentor_uuid, 'school_ids.json'), 'r') as f:
            school_ids = json.load(f)
        school_name = school_ids["uuid_school"][school_uuid]
        with open(os.path.join('assets', mentor_uuid, school_name, 'metadata.json'), 'r') as f:
            metadata = json.load(f)

        student_num = metadata['uuid_student_num'][student_uuid]
        file_name = student_num + '_report.json'
        if os.path.exists(os.path.join('assets', mentor_uuid, school_name, file_name)):
            with open(os.path.join('assets', mentor_uuid, school_name, file_name), 'r') as f:
                saved_data = json.load(f)
        else:
            saved_data = {"contents": [], "eval": '', "eval_revised": ""}
        return jsonify(saved_data)

    else:
        data = request.json
        mentor_uuid = request.args['mentor_code']
        school_uuid = request.args['school_code']
        student_uuid = request.args['student_code']
        with open(os.path.join('assets', mentor_uuid, 'school_ids.json'), 'r') as f:
            school_ids = json.load(f)
        school_name = school_ids["uuid_school"][school_uuid]
        with open(os.path.join('assets', mentor_uuid, school_name, 'metadata.json'), 'r') as f:
            metadata = json.load(f)

        student_num = metadata['uuid_student_num'][student_uuid]
        file_name = student_num + '_report.json'

        if os.path.exists(os.path.join('assets', mentor_uuid, school_name, file_name)):
            with open(os.path.join('assets', mentor_uuid, school_name, file_name), 'r') as f:
                save_data = json.load(f)
        else:
            save_data = {"contents": [], "eval": '', "eval_revised": {"text": "", "class_attitude": "중", "class_attitude_detail": "", "report_topic" : ""}}

        if data['type'] == 'report':
            save_data['contents'] = data['contents']
            metadata['report_status'][student_num]['final_report'] = 'Yes'
        elif data['type'] == 'eval':
            save_data['eval'] = data['content']
            metadata['report_status'][student_num]['self_evaluation'] = 'Yes'
        elif data['type'] == 'eval_revised':
            save_data['eval_revised'] = data['content']
            metadata['report_status'][student_num]['self_evaluation_revised'] = 'Yes'

        with open(os.path.join('assets', mentor_uuid, school_name, file_name), 'w') as f:
            json.dump(save_data, f)

        with open(os.path.join('assets', mentor_uuid, school_name, 'metadata.json'), 'w') as f:
            json.dump(metadata, f)

        return jsonify({"message" : "성공적으로 저장했습니다"})

@app.route('/delete_school/<mentor_uuid>', methods=['POST'])
def delete_school(mentor_uuid):
    data = request.json
    school_name = data['school_name']
    with open(os.path.join('assets', mentor_uuid, 'school_ids.json'), 'r') as f:
        school_ids = json.load(f)
    school_id = school_ids['school_uuid'][school_name]
    school_ids['school_uuid'].pop(school_name)
    school_ids['uuid_school'].pop(school_id)
    shutil.rmtree(os.path.join('assets', mentor_uuid, school_name))
    return jsonify({"message" : "Success"})

@app.route('/final_report_modify')
def final_report_modify():
    return render_template('main_mentor.html')

@app.route('/total_report/<report_name>', methods=['GET', 'POST'])
def get_total_report(report_name):
    mentor_uuid = request.args['mentor_code']
    school_uuid = request.args['school_code']
    with open(os.path.join('assets', mentor_uuid, 'school_ids.json'), 'r') as f:
        school_ids = json.load(f)
    school_name = school_ids["uuid_school"][school_uuid]

    if request.method == 'POST':
        data = request.json
        with open(os.path.join('assets', mentor_uuid, school_name, 'mentor_' + report_name + '.json'), 'w') as f:
            json.dump(data, f)
        return jsonify({"message" : "성공적으로 저장되었습니다!"})

    elif request.method == 'GET':
        if report_name == 'dropdown':
            files = os.listdir(os.path.join('assets', mentor_uuid, school_name))
            res = []
            for file in files:
                if 'mentor_' in file:
                    res.append(file.replace('mentor_', '').replace('.json', ''))
            return jsonify({"message": "Success", "files": res})

        elif not os.path.exists(os.path.join('assets', mentor_uuid, school_name, 'mentor_' + report_name + '.json')):
            return jsonify({"message": "해당 파일이 없습니다"})

        else:
            with open(os.path.join('assets', mentor_uuid, school_name, 'mentor_' + report_name + '.json'), 'r') as f:
                res = json.load(f)
            return jsonify({"message": "Success", "data": res})

@app.route('/export_total_report', methods=['POST'])
def export_total_report():
    global counters, table_counter, image_counter, ref_counter, previous_level
    mentor_uuid = request.args['mentor_code']
    school_uuid = request.args['school_code']
    report_name = request.json['report_name']
    file_name = request.json['file_name']
    with open(os.path.join('assets', mentor_uuid, 'school_ids.json'), 'r') as f:
        school_ids = json.load(f)
    school_name = school_ids["uuid_school"][school_uuid]
    with open(os.path.join('assets', mentor_uuid, school_name, 'metadata.json'), 'r') as f:
        metadata = json.load(f)
    with open(os.path.join('assets', mentor_uuid, school_name, 'mentor_' + report_name + '.json'), 'r') as f:
        report = json.load(f)

    doc = Document('FinalReport_empty.docx')
    start_date = report['metadata']['startDate']
    end_date = report['metadata']['endDate']
    doc.tables[0].rows[0].cells[2].paragraphs[0].text = metadata['school_name']
    doc.tables[0].rows[0].cells[4].paragraphs[0].text = report['metadata']['students']
    doc.tables[0].rows[1].cells[2].paragraphs[0].text = metadata['field']
    doc.tables[0].rows[1].cells[4].paragraphs[0].text = metadata['class_info']['mentorName']
    doc.tables[0].rows[2].cells[1].paragraphs[0].text = f"{report['metadata']['startDate'][0:4]}년 {start_date[4:6]}월 {start_date[6:]}일 ~ {end_date[0:4]}년 {end_date[4:6]}월 {end_date[6:]}일"
    doc.tables[0].rows[3].cells[2].paragraphs[0].text = report['metadata']['reportName']
    doc.tables[0].rows[4].cells[2].paragraphs[0].text = report['metadata']['reportNameEng']
    doc.tables[0].rows[5].cells[1].paragraphs[0].text = report['metadata']['keyWords']

    for content in report['contents']:
        if content['type'] == 'h1':
            current_level = int(content['type'][1]) - 1
            if current_level < previous_level:
                for i in range(current_level + 1, 7):
                    counters[i] = i % 2

            para = doc.tables[1].rows[0].cells[0].add_paragraph()
            run = para.add_run(h1_headers[counters[current_level]] + content['content'])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '맑은 고딕'
            counters[current_level] += 1
            previous_level = current_level

        elif content['type'] == 'text':
            para = doc.tables[1].rows[0].cells[0].add_paragraph()
            run = para.add_run(' ' + content['content'] + '\n')
            run.bold = False
            run.font.size = Pt(10)

        elif content['type'] == 'image':
            para = doc.tables[1].rows[0].cells[0].add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            binary = base64.b64decode(content['content'].split(',')[1])
            run.add_picture(io.BytesIO(binary), height=Inches(3))
            run2 = para.add_run('\n그림 ' + str(image_counter) + '. ' + content['caption'] + '\n')
            run2.font.size = Pt(8)
            image_counter += 1

        elif content['type'] == "ref":
            para = doc.tables[1].rows[0].cells[0].add_paragraph()
            run = para.add_run(f'[{ref_counter}] ' + content['content'])
            run.bold = False
            run.font.size = Pt(10)
            ref_counter += 1

        else:
            current_level = int(content['type'][1]) - 1
            if current_level < previous_level:
                for i in range(current_level + 1, 7):
                    counters[i] = i % 2

            para = doc.tables[1].rows[0].cells[0].add_paragraph()
            para.paragraph_format.left_indent = Pt(20 * current_level)
            run = para.add_run(get_current_header(current_level, counters[current_level]) + content['content'])
            run.bold = False
            run.font.size = Pt(10)
            counters[current_level] += 1
            previous_level = current_level

    image_counter = 1
    ref_counter = 1
    counters = [0, 1, 0, 1, 0, 1, 0]
    previous_level = 0
    final_dir = os.path.join('assets', mentor_uuid, school_name, 'exported_' + report_name + '.docx')
    doc.save(final_dir)
    return send_file(final_dir, download_name=file_name + '.docx', as_attachment=True)

@app.route('/export_self_eval', methods=['POST'])
def export_self_eval():
    mentor_uuid = request.args['mentor_code']
    school_uuid = request.args['school_code']
    student_name = request.json['student_name']
    student_num = request.json['student_num']
    with open(os.path.join('assets', mentor_uuid, 'school_ids.json'), 'r') as f:
        school_ids = json.load(f)
    school_name = school_ids["uuid_school"][school_uuid]
    with open(os.path.join('assets', mentor_uuid, school_name, 'metadata.json'), 'r') as f:
        metadata = json.load(f)
    with open(os.path.join('assets', mentor_uuid, school_name, student_num + '_report.json'), 'r') as f:
        report = json.load(f)

    doc = Document('SelfEval_empty.docx')
    doc.tables[0].rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.tables[0].rows[0].cells[1].paragraphs[0].text = metadata['school_name']
    doc.tables[0].rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.tables[0].rows[1].cells[1].paragraphs[0].text = student_num + ' ' + student_name
    doc.tables[0].rows[2].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.tables[0].rows[2].cells[1].paragraphs[0].text = metadata['class_info']['mentorName']
    doc.tables[0].rows[3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.tables[0].rows[3].cells[1].paragraphs[0].text = report['eval_revised']['report_topic']
    doc.tables[1].rows[0].cells[0].paragraphs[0].text = report['eval']
    doc.tables[2].rows[1].cells[0].paragraphs[0].text = report['eval_revised']['class_attitude']
    doc.tables[2].rows[1].cells[1].paragraphs[0].text = report['eval_revised']['class_attitude_detail']
    doc.tables[2].rows[2].cells[0].paragraphs[0].text = report['eval_revised']['text']
    final_dir = os.path.join('assets', mentor_uuid, school_name, '자기활동평가서_' + student_name + '.docx')
    doc.save(final_dir)
    return send_file(final_dir, download_name='자기활동평가서_' + student_name + '.docx', as_attachment=True)

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5002, debug=False)